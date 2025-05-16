import os
import sys
import re

try:
    import docx
except ImportError:
    print("python-docx not installed. Installing...")
    os.system(f"{sys.executable} -m pip install python-docx")
    import docx

try:
    import markdown
except ImportError:
    print("markdown not installed. Installing...")
    os.system(f"{sys.executable} -m pip install markdown")
    import markdown

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_table(text):
    """Extract markdown table from text"""
    lines = text.strip().split('\n')
    table_data = []
    header = []
    
    if len(lines) < 3:
        return None, text
    
    # Check if this is a table
    if not lines[0].startswith('|') or not lines[1].startswith('|'):
        return None, text
    
    # Find the end of the table
    table_end = 0
    for i, line in enumerate(lines):
        if not line.startswith('|'):
            table_end = i
            break
        if i == len(lines) - 1:
            table_end = i + 1
    
    if table_end < 3:  # Not a valid table
        return None, text
    
    # Process header
    header_parts = lines[0].strip().split('|')
    header = [part.strip() for part in header_parts if part.strip()]
    
    # Process rows
    for i in range(2, table_end):  # Skip separator line
        row_parts = lines[i].strip().split('|')
        row = [part.strip() for part in row_parts if part or part == '']
        if row:
            table_data.append(row)
    
    remaining_text = '\n'.join(lines[table_end:])
    
    return (header, table_data), remaining_text

def add_table_to_doc(doc, table_data):
    """Add a table to the Word document"""
    if not table_data:
        return
    
    header, rows = table_data
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = 'Table Grid'
    
    # Add header
    for i, text in enumerate(header):
        cell = table.cell(0, i)
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add rows
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            if j < len(header):  # Ensure we don't exceed the number of columns
                cell = table.cell(i + 1, j)
                cell.text = text
    
    # Add a paragraph after the table
    doc.add_paragraph()

def process_text_with_tables(doc, text):
    """Process text and extract/add tables"""
    remaining = text
    
    while remaining:
        table_data, remaining = extract_table(remaining)
        
        if table_data:
            # Add text before table
            parts = remaining.split('\n', 1)
            if parts and parts[0].strip():
                doc.add_paragraph(parts[0].strip())
            if len(parts) > 1:
                remaining = parts[1]
            else:
                remaining = ""
            
            # Add the table
            add_table_to_doc(doc, table_data)
        else:
            # No more tables, add remaining text
            if remaining.strip():
                paragraphs = remaining.strip().split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            remaining = ""

def md_to_docx(md_file, docx_file):
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Set font size for normal text
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    
    # Simple conversion - add paragraphs
    sections = md_content.split('\n## ')
    
    # Handle the first section (title and first section)
    first_parts = sections[0].split('\n# ', 1)
    if len(first_parts) > 1:
        # Add the main title
        title = first_parts[1].strip()
        doc.add_heading(title, level=0)
        
        # Add the first section content if any
        if '\n' in first_parts[1]:
            content = '\n'.join(first_parts[1].split('\n')[1:])
            if content.strip():
                process_text_with_tables(doc, content.strip())
    
    # Process the remaining sections
    for section in sections[1:]:
        section_parts = section.split('\n', 1)
        if section_parts:
            # Add section title
            doc.add_heading(section_parts[0], level=1)
            
            # Add section content if any
            if len(section_parts) > 1:
                # Split into subsections
                subsections = section_parts[1].split('\n### ')
                
                # Add main section content
                if subsections[0].strip():
                    process_text_with_tables(doc, subsections[0].strip())
                
                # Process subsections
                for subsection in subsections[1:]:
                    sub_parts = subsection.split('\n', 1)
                    if sub_parts:
                        # Add subsection title
                        doc.add_heading(sub_parts[0], level=2)
                        
                        # Add subsection content
                        if len(sub_parts) > 1 and sub_parts[1].strip():
                            process_text_with_tables(doc, sub_parts[1].strip())
    
    # Save the document
    doc.save(docx_file)
    print(f"Converted {md_file} to {docx_file}")

if __name__ == "__main__":
    md_file = "pwv数据分析结果.md"
    docx_file = "PWV数据分析结果报告.docx"
    md_to_docx(md_file, docx_file) 